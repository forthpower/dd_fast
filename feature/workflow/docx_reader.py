#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Docx文件读取器
从Primer Workflow文档中提取分量配置信息
"""

import re
from typing import Dict, List, Optional
try:
    from docx import Document
except ImportError:
    Document = None


class DocxReader:
    """Docx文件读取器"""
    
    def __init__(self):
        if Document is None:
            raise ImportError("请安装python-docx库: pip install python-docx")
    
    def read_docx(self, file_path: str) -> Dict[str, Dict[str, List[int]]]:
        """
        读取docx文件并提取分量配置
        
        Args:
            file_path: docx文件路径
            
        Returns:
            分量配置字典: {支付方式: {币种: [Adyen%, Stripe%, AWX%]}}
        """
        doc = Document(file_path)
        result = {}
        current_payment_method = None
        
        # 遍历所有段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 检测支付方式（CARD, AP, GP等）
            payment_match = re.match(r'^([A-Z]+)$', text)
            if payment_match:
                current_payment_method = payment_match.group(1)
                result[current_payment_method] = {}
                continue
            
            # 检测币种和分量配置
            # 格式: 币种 - 20%：40%：40% 或 币种1/币种2/币种3 - 20%：40%：40%
            if current_payment_method:
                pattern = r'^([A-Z/]+)\s*-\s*(\d+)%[：:]\s*(\d+)%[：:]\s*(\d+)%'
                match = re.match(pattern, text)
                if match:
                    currencies_str = match.group(1)
                    adyen_pct = int(match.group(2))
                    stripe_pct = int(match.group(3))
                    awx_pct = int(match.group(4))
                    
                    # 分割多个币种
                    currencies = [c.strip() for c in currencies_str.split('/')]
                    
                    for currency in currencies:
                        if currency not in result[current_payment_method]:
                            result[current_payment_method][currency] = []
                        result[current_payment_method][currency] = [adyen_pct, stripe_pct, awx_pct]
        
        # 也检查表格
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join([cell.text.strip() for cell in row.cells])
                if not row_text:
                    continue
                
                # 检测支付方式
                payment_match = re.match(r'^([A-Z]+)$', row_text)
                if payment_match:
                    current_payment_method = payment_match.group(1)
                    if current_payment_method not in result:
                        result[current_payment_method] = {}
                    continue
                
                # 检测币种和分量配置
                if current_payment_method:
                    pattern = r'^([A-Z/]+)\s*-\s*(\d+)%[：:]\s*(\d+)%[：:]\s*(\d+)%'
                    match = re.match(pattern, row_text)
                    if match:
                        currencies_str = match.group(1)
                        adyen_pct = int(match.group(2))
                        stripe_pct = int(match.group(3))
                        awx_pct = int(match.group(4))
                        
                        currencies = [c.strip() for c in currencies_str.split('/')]
                        
                        for currency in currencies:
                            if currency not in result[current_payment_method]:
                                result[current_payment_method][currency] = []
                            result[current_payment_method][currency] = [adyen_pct, stripe_pct, awx_pct]
        
        return result

